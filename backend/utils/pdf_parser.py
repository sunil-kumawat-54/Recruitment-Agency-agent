import re
import os
import PyPDF2
from docx import Document
from io import BytesIO

def extract_text_from_pdf(file_input) -> str:
    """Extracts raw text strings from PDF files using PyPDF2 PdfReader."""
    text = ""
    try:
        # Check if input is a path string or a binary byte stream
        if isinstance(file_input, str):
            with open(file_input, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        else:
            # Handle UploadFile memory byte streams
            reader = PyPDF2.PdfReader(BytesIO(file_input))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[PDF Extraction Error] failed to parse pdf: {e}")
    return text

def extract_text_from_docx(file_input) -> str:
    """Extracts raw text strings from DOCX files using python-docx Document."""
    text = ""
    try:
        if isinstance(file_input, str):
            doc = Document(file_input)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            doc = Document(BytesIO(file_input))
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        print(f"[DOCX Extraction Error] failed to parse docx: {e}")
    return text

def parse_resume_to_dict(file_input, filename: str) -> dict:
    """Parses a resume file and extracts standard sections, contact info, and word counts."""
    file_type = "txt"
    full_text = ""
    
    # Identify file format from extension
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        file_type = "pdf"
        full_text = extract_text_from_pdf(file_input)
    elif ext in ["docx", "doc"]:
        file_type = "docx"
        full_text = extract_text_from_docx(file_input)
    else:
        # Fallback to plain string conversion
        if isinstance(file_input, bytes):
            full_text = file_input.decode("utf-8", errors="ignore")
        else:
            full_text = str(file_input)
            
    word_count = len(full_text.split())
    
    # Categorize document sections based on header markers
    sections = {
        "contact_info": "",
        "skills": "",
        "experience": "",
        "education": "",
        "certifications": ""
    }
    
    # Basic structural partitioning heuristics
    lines = full_text.split("\n")
    current_section = "contact_info"
    
    section_patterns = {
        "skills": re.compile(r'\b(skills|competencies|technologies|proficiencies)\b', re.I),
        "experience": re.compile(r'\b(experience|employment|work history|career history)\b', re.I),
        "education": re.compile(r'\b(education|academic|university|qualifications)\b', re.I),
        "certifications": re.compile(r'\b(certifications|credentials|awards|licenses)\b', re.I)
    }
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        # Check if line indicates a header section transition
        transitioned = False
        for sec, pattern in section_patterns.items():
            # If line is short and matches a keyword, switch section
            if len(line_stripped) < 40 and pattern.search(line_stripped):
                current_section = sec
                transitioned = True
                break
                
        if transitioned:
            continue
            
        sections[current_section] += line_stripped + "\n"

    # Quick win: Extract basic contact info (emails / phones) if contact section is empty
    if not sections["contact_info"].strip():
        emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', full_text)
        phones = re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', full_text)
        sections["contact_info"] = f"Emails: {', '.join(emails)} | Phones: {', '.join(phones)}"
        
    return {
        "full_text": full_text,
        "sections": sections,
        "word_count": word_count,
        "file_type": file_type
    }
